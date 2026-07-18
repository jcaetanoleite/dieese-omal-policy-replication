from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from zipfile import ZipFile, ZIP_DEFLATED
import shutil
import re

ROOT = Path(__file__).resolve().parents[1]
FIG_DIR = ROOT / 'outputs' / 'figures'
ARTICLE_DIR = ROOT / 'outputs' / 'article'
TABLE_DIR = ROOT / 'outputs' / 'tables'
ARTICLE_DIR.mkdir(parents=True, exist_ok=True)
OUT = ARTICLE_DIR / 'Is_the_DIEESE_minimum_wage_useful_for_policy_editorial_v2.docx'
MD_OUT = ARTICLE_DIR / 'Is_the_DIEESE_minimum_wage_useful_for_policy_editorial_v2.md'
ZIP_OUT = ARTICLE_DIR / 'DIEESE_policy_article_editorial_v2.zip'
PKG = ARTICLE_DIR / 'package'

# Figure path mapping, preserving the original figures 1-14 and high-resolution OMAL figures 15-22.
figures = {
    1: FIG_DIR/'figure01.png',
    2: FIG_DIR/'figure02.png',
    3: FIG_DIR/'figure03.png',
    4: FIG_DIR/'figure04.png',
    5: FIG_DIR/'figure05.png',
    6: FIG_DIR/'figure06.png',
    7: FIG_DIR/'figure07.png',
    8: FIG_DIR/'figure08.png',
    9: FIG_DIR/'figure09.png',
    10: FIG_DIR/'figure10.png',
    11: FIG_DIR/'figure11.png',
    12: FIG_DIR/'figure12.png',
    13: FIG_DIR/'figure13.png',
    14: FIG_DIR/'figure14.png',
    15: FIG_DIR/'figure15.png',
    16: FIG_DIR/'figure16.png',
    17: FIG_DIR/'figure17.png',
    18: FIG_DIR/'figure18.png',
    19: FIG_DIR/'figure19.png',
    20: FIG_DIR/'figure20.png',
    21: FIG_DIR/'figure21.png',
    22: FIG_DIR/'figure22.png',
}
for n,p in figures.items():
    if not p.exists():
        raise FileNotFoundError(f'Missing figure {n}: {p}')

# ---------------- Document setup ----------------
doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.72)
sec.bottom_margin = Inches(0.72)
sec.left_margin = Inches(0.82)
sec.right_margin = Inches(0.82)

styles = doc.styles
normal = styles['Normal']
normal.font.name = 'Georgia'
normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Georgia')
normal.font.size = Pt(10.7)
normal.paragraph_format.space_after = Pt(7)
normal.paragraph_format.line_spacing = 1.13
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

for style_name, size in [('Title', 25), ('Subtitle', 11.5), ('Heading 1', 15), ('Heading 2', 12.5)]:
    st = styles[style_name]
    st.font.name = 'Georgia'
    st._element.rPr.rFonts.set(qn('w:eastAsia'), 'Georgia')
    st.font.size = Pt(size)
    if style_name.startswith('Heading'):
        st.font.bold = True
        st.font.color.rgb = RGBColor(25,25,25)
        st.paragraph_format.space_before = Pt(15)
        st.paragraph_format.space_after = Pt(7)

# Custom styles
if 'Deck' not in styles:
    deck = styles.add_style('Deck', WD_STYLE_TYPE.PARAGRAPH)
else:
    deck = styles['Deck']
deck.font.name = 'Georgia'
deck._element.rPr.rFonts.set(qn('w:eastAsia'), 'Georgia')
deck.font.size = Pt(11.2)
deck.font.italic = True
deck.font.color.rgb = RGBColor(70,70,70)
deck.paragraph_format.space_after = Pt(14)

if 'Figure Caption' not in styles:
    capstyle = styles.add_style('Figure Caption', WD_STYLE_TYPE.PARAGRAPH)
else:
    capstyle = styles['Figure Caption']
capstyle.font.name = 'Georgia'
capstyle._element.rPr.rFonts.set(qn('w:eastAsia'), 'Georgia')
capstyle.font.size = Pt(9.1)
capstyle.font.italic = True
capstyle.font.color.rgb = RGBColor(45,45,45)
capstyle.paragraph_format.space_after = Pt(2)

if 'Figure Source' not in styles:
    srcstyle = styles.add_style('Figure Source', WD_STYLE_TYPE.PARAGRAPH)
else:
    srcstyle = styles['Figure Source']
srcstyle.font.name = 'Georgia'
srcstyle._element.rPr.rFonts.set(qn('w:eastAsia'), 'Georgia')
srcstyle.font.size = Pt(7.8)
srcstyle.font.color.rgb = RGBColor(100,100,100)
srcstyle.paragraph_format.space_after = Pt(8)

if 'Equation' not in styles:
    eqstyle = styles.add_style('Equation', WD_STYLE_TYPE.PARAGRAPH)
else:
    eqstyle = styles['Equation']
eqstyle.font.name = 'Cambria Math'
eqstyle._element.rPr.rFonts.set(qn('w:eastAsia'), 'Cambria Math')
eqstyle.font.size = Pt(10.4)
eqstyle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
eqstyle.paragraph_format.space_before = Pt(5)
eqstyle.paragraph_format.space_after = Pt(7)

if 'Reference' not in styles:
    refstyle = styles.add_style('Reference', WD_STYLE_TYPE.PARAGRAPH)
else:
    refstyle = styles['Reference']
refstyle.font.name = 'Georgia'
refstyle._element.rPr.rFonts.set(qn('w:eastAsia'), 'Georgia')
refstyle.font.size = Pt(8.7)
refstyle.paragraph_format.left_indent = Inches(0.25)
refstyle.paragraph_format.first_line_indent = Inches(-0.25)
refstyle.paragraph_format.space_after = Pt(4)

# Footer page number field
footer_p = sec.footer.paragraphs[0]
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer_p.add_run()
fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve'); instrText.text = ' PAGE '
fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end')
run._r.extend([fldChar1, instrText, fldChar2])

# ---------------- Helpers ----------------
def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def add_para(text, italic=False, bold_start=None):
    p = doc.add_paragraph(style='Normal')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_start and text.startswith(bold_start):
        r1 = p.add_run(bold_start)
        r1.bold = True
        p.add_run(text[len(bold_start):])
    else:
        r = p.add_run(text)
        r.italic = italic
    return p


def add_equation(text):
    p = doc.add_paragraph(style='Equation')
    p.add_run(text)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    return p


def add_figure(n, caption, source, width=6.35):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(figures[n]), width=Inches(width))
    cp = doc.add_paragraph(f'Figure {n}. {caption}', style='Figure Caption')
    cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cp.paragraph_format.keep_with_next = True
    sp = doc.add_paragraph(f'Source: {source}', style='Figure Source')
    sp.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_heading(text, level=1):
    return doc.add_heading(text, level=level)

# ---------------- Article ----------------
title = doc.add_paragraph(style='Title')
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
title.add_run('Is the DIEESE minimum wage useful for policy?')

deck = doc.add_paragraph(style='Deck')
deck.add_run('DIEESE gave Brazilian workers an independent economic vocabulary. Its “necessary minimum wage” remains useful in bargaining, but the formula is too rigid to measure household need and too detached from labour-market evidence to guide the statutory wage floor.')

add_para(
    "Each month, DIEESE publishes two numbers. The first is the cost of a basic food basket. The second, and more politically resonant, is the income that a four-person household would supposedly need to cover the goods and services named in Brazil's minimum-wage legislation. In June 2026 the second figure was R$8,110.92, nearly five times the statutory minimum of R$1,621. The contrast is cited by unions, newspapers and progressive commentators as a compact account of the distance between the legal wage and a decent standard of living (DIEESE, 2026; Ipea, 2026)."
)
add_para(
    "The appeal is understandable. A social right written in constitutional language is difficult to take into a bargaining room. A number is easier. It can be compared with a wage offer, printed on a placard and followed over time. Yet a number can be politically effective and empirically weak at the same time. The DIEESE benchmark works best as a statement of distributive ambition. It becomes much harder to defend when it is treated as an estimate of what a representative household needs, as a poverty threshold or as a target for the statutory minimum wage."
)
add_para(
    "The distinction matters because DIEESE is not an incidental institution. It was created in 1955 by trade unions that wanted their own evidence on prices, wages and employment. Over the following decades it measured inflation, supported collective bargaining and supplied organised labour with technical expertise that employers and the state had long possessed on better terms (DIEESE, 2025). The necessary-minimum-wage series belongs to that history. Criticising the formula does not diminish the institution. It asks the number to carry only the weight that its method can bear."
)

add_figure(1, "The real value of Brazil's minimum wage, 1940-2026.", "author's calculations from the Ipea real minimum-wage series (Ipea, 2026).")

add_para(
    "The legal minimum wage has itself been transformed over that history. Its real value fell sharply before the stabilisation of the 1990s, rose for much of the 2000s and 2010s, and reached R$1,621 in June 2026 in the price base used in Figure 1 (Ipea, 2026). Those movements affected much more than the pay of workers exactly at the floor. The minimum shapes formal wage contracts, spills into some informal wages, anchors parts of the wage distribution and indexes pensions and social benefits. It also changes firms' labour costs and public expenditure. A policy rule for such an institution must therefore be connected to evidence on employment, hours, compliance, prices, productivity, informality and fiscal exposure (Dube and Lindner, 2024; Instituição Fiscal Independente, 2026)."
)

add_heading('The formula')

add_para(
    "DIEESE begins with the most expensive basic food basket among the capitals in its monthly survey. The quantities descend from the regional baskets associated with Decree-Law 399 of 1938. The cost is multiplied by three because two adults and two children are treated as three adult food consumers. The resulting food bill is divided by 0.3571, the food share estimated in DIEESE's own 1994-95 household-budget survey (DIEESE, 2016)."
)
add_equation(r"SMN_t = \frac{3C_t^{max}}{0.3571}")
add_para(
    "Here, C_t^{max} is the price of the most expensive basket in month t. Housing, transport, health, education, clothing and leisure are not priced directly. Their combined cost is the residual implied by the claim that food represents 35.71 per cent of the household budget. The formula is admirably transparent. Its transparency also reveals the problem: almost two-thirds of the final estimate is generated by a fixed coefficient rather than by current observations of non-food spending (DIEESE, 2016)."
)
add_para(
    "The age of the 1938 basket is not, by itself, fatal. Old standards can remain useful when their nutritional content, quantities and reference population are repeatedly tested against new evidence. In this case the basket's legal ancestry often substitutes for that validation. The coefficient in the denominator creates an even larger difficulty. Brazil has since conducted national household-budget surveys in 2002-03, 2008-09 and 2017-18, and each one records large differences in spending patterns across income groups and over time (IBGE, 2004, 2010, 2019)."
)

add_figure(2, "Food expenditure shares across four household-budget surveys.", "Ipea tabulations of the 1995-96 survey and POF releases (IBGE, 2004, 2010, 2019). Income classes and geographic coverage are not identical across surveys.")

add_para(
    "Figure 2 illustrates Engel's law. Poorer households devote a much larger share of consumption to food, and that share falls as income rises. The composition of non-food expenditure has also changed as housing, transport, health, education and communications have taken a larger place in family budgets (IBGE, 2004, 2010, 2019). A coefficient estimated in 1994-95 cannot represent every income group, region and year without an argument for why household expenditure should have remained structurally unchanged."
)
add_para(
    "Replacing 35.71 per cent with the food share from the latest POF would not solve the problem. Because the coefficient sits in the denominator, a lower food share would mechanically raise the total budget. The difficult work lies elsewhere: choosing the households that define adequacy, deciding how family size changes needs, pricing large non-food items directly and allowing for regional variation. These are not technical details that can be delegated to a single coefficient. As one review of poverty-line methods puts it, 'todos os problemas de escolhas de variáveis, pesos e uma linha de corte dependem de julgamentos de valor bastante arbitrários' [the choices of variables, weights and a cutoff depend on fairly arbitrary value judgments] (Soares, 2009, p. 15)."
)
add_para(
    "The household rule is equally blunt. The assumption that two children consume as much food as one adult is carried from the food basket into a total budget containing school costs, childcare, clothing, health care and housing space. Those items do not scale with age in the same way. Modern equivalence scales are also normative, but they make the norm visible and allow sensitivity analysis. The DIEESE method fixes one family type and one equivalence rule inside the headline number (DIEESE, 2016; OECD, 2013)."
)
add_para(
    "The use of the most expensive capital adds a third layer. The series is not a population-weighted average, a median cost or a set of regional thresholds. It follows the upper edge of the surveyed cities. That can be an effective bargaining strategy. It is a poor description of the household budget faced by a typical family in a country with large differences in prices, earnings and public provision."
)

add_figure(3, "The statutory minimum wage and the DIEESE benchmark per person, 1994-2026.", "DIEESE and Ipea; dividing the family benchmark by four is a sensitivity exercise, not part of the DIEESE method (DIEESE, 2026; Ipea, 2026).")
add_para(
    "Dividing the family estimate by four does not repair the conceptual mismatch, but it places the two magnitudes on a more intelligible scale. For much of the period the per-person DIEESE figure was close to the legal minimum. It moved above it when food inflation accelerated after 2020. That movement is real, but its interpretation is narrower than the public discussion often suggests. It shows food prices passing through a fixed family and budget coefficient; it does not show that every other household cost moved in the same proportion (DIEESE, 2016, 2026)."
)

add_figure(4, "The DIEESE benchmark per person as a share of the statutory minimum, 1994-2026.", "author's calculations from DIEESE and Ipea series (DIEESE, 2026; Ipea, 2026).")

add_figure(5, "State household income per capita relative to the per-person DIEESE benchmark, 2025.", "PNAD Contínua and DIEESE (IBGE, 2026a; DIEESE, 2026).", width=5.85)
add_para(
    "The cross-section in Figure 5 makes the regional problem visible. The Federal District and several states in the South and Southeast sit well above the per-person benchmark. Many states in the North and Northeast sit below it. This does not prove that the benchmark is accurate in low-income states; household income is not a local price index. It does show that one national number is being compared with labour markets that differ sharply in productivity, employment structure and earnings (IBGE, 2026a)."
)
add_para(
    "The quarterly panels use a different concept: average habitual income from all jobs among employed people. That series is available consistently by state from 2012, whereas household income per capita is not available in the same quarterly format. The distinction matters, and the figures state it explicitly (IBGE, 2026a)."
)

add_figure(6, "North: quarterly real labour income and the DIEESE benchmark, 2012-2026.", "author's calculations from PNAD Contínua, DIEESE and Ipea (IBGE, 2026a; DIEESE, 2026; Ipea, 2026).")
add_figure(7, "Northeast: quarterly real labour income and the DIEESE benchmark, 2012-2026.", "author's calculations from PNAD Contínua, DIEESE and Ipea (IBGE, 2026a; DIEESE, 2026; Ipea, 2026).")
add_figure(8, "Southeast: quarterly real labour income and the DIEESE benchmark, 2012-2026.", "author's calculations from PNAD Contínua, DIEESE and Ipea (IBGE, 2026a; DIEESE, 2026; Ipea, 2026).")
add_figure(9, "South: quarterly real labour income and the DIEESE benchmark, 2012-2026.", "author's calculations from PNAD Contínua, DIEESE and Ipea (IBGE, 2026a; DIEESE, 2026; Ipea, 2026).")
add_figure(10, "Center-West: quarterly real labour income and the DIEESE benchmark, 2012-2026.", "author's calculations from PNAD Contínua, DIEESE and Ipea (IBGE, 2026a; DIEESE, 2026; Ipea, 2026).")

add_para(
    "Across these panels, the per-person benchmark generally remains below average labour income, but the distance is much narrower in the Northeast and in parts of the North. The Federal District, São Paulo and the southern states show much wider gaps. The figures do not establish where the legal minimum should be set. They show why a single family-cost estimate is a fragile guide to a labour market whose regional structures differ so much (IBGE, 2026a; DIEESE, 2026)."
)

add_figure(11, "Informality and the ratio of average labour income to the DIEESE benchmark per person, by region, 2015Q4-2026Q1.", "state-quarter observations from PNAD Contínua and the DIEESE benchmark. The fitted lines are descriptive OLS regressions (IBGE, 2026a, 2026c; DIEESE, 2026).")
add_para(
    "Figure 11 relates the earnings gap to informality. The association is negative in four regions, strongest in the Center-West and Southeast, milder in the North and Northeast, and nearly flat in the South. It should not be read causally. Informality, productivity, urbanisation, sectoral composition and wages are jointly determined. The pattern is nevertheless consistent with a practical concern: a national budget benchmark ignores the structure of the local labour market in which a wage floor must operate (IBGE, 2026a, 2026c; Ulyssea, 2020)."
)

add_heading('The wage floor is a different object')

add_para(
    "A statutory minimum wage regulates an employment relationship. The DIEESE estimate describes a four-person household with one implicit earner and no other income. Brazilian households differ in size, the number of workers, hours supplied, benefit receipt and access to public services. Converting one household type into a wage target would write a particular family model into labour policy without demonstrating that it is representative."
)
add_para(
    "The scale of the proposal is no less important. In the broad PNAD comparison, the full DIEESE estimate exceeded 350 per cent of median labour income in 2023. A rapid attempt to close a gap of that size would affect formal employment, informal work, hours, prices, firm entry and exit, pensions, social assistance and public budgets. None of those margins appears in the household-cost formula."
)
add_para(
    "The modern minimum-wage literature is useful precisely because it does not reduce the question to one employment coefficient. A major review finds that minimum wages have often produced 'limited direct employment effects while significantly increasing the earnings of low-wage workers', but also documents responses in prices, profits, turnover, productivity, amenities and capital use (Dube and Lindner, 2024, p. 67). Those findings support neither a reflexive fear of wage floors nor an unlimited one. They support policy that attends to the level of the floor and the institutions around it."
)
add_para(
    "Individual trajectories can diverge even when the aggregate result looks modest. Evidence from Seattle found 'some workers enjoying increased earnings and continued employment opportunities while others see a decline in their prospects' (Jardim et al., 2018, p. 3). A policy benchmark therefore needs to say something about who gains, who loses and how adjustment occurs. A formula based on a food basket cannot answer those questions."
)
add_para(
    "The criticism should not obscure the political value of the DIEESE number. Trade unions bargain in markets where firms often possess more information and more power. When rents are present, 'institutional changes that enhance workers' countervailing power' may raise labour's share (Stansbury and Summers, 2020, p. 9). An independent cost-of-living benchmark can strengthen that countervailing power even when it is unsuitable as an administrative rule."
)
add_para(
    "Brazilian evidence also shows why the legal minimum deserves careful analysis on its own terms. The large increases between 2000 and 2009 passed through almost fully to informal employees in formal firms and more partially to informal workers in informal firms. The minimum can therefore 'positively affect living standards for workers thought beyond the reach of labor law' (Derenoncourt et al., 2025, p. 2). The result strengthens the case for the statutory floor while weakening the case for replacing labour-market evidence with a family-cost formula."
)
add_para(
    "The distributional effects extended above the floor. The 128 per cent real increase between 1996 and 2018 generated wage spillovers and accounted for '45 percent of a large fall in earnings inequality' (Engbom and Moser, 2022, p. 1). Limited employment and output effects were partly absorbed through worker reallocation towards more productive firms. Internal wage structures also moved with the floor: firms appear to 'anchor their internal wage structure on the minimum wage due to fairness or internal incentives issues' (Brochu et al., 2025, p. 2). Other responses appear in hours, job search, automation and workplace safety (Jardim et al., 2018; Melo et al., 2026; Brynjolfsson et al., 2026; Davies et al., 2026)."
)
add_para(
    "Informality makes the Brazilian case still more demanding. It is produced by decisions of heterogeneous firms and workers under different levels of productivity, enforcement and regulatory cost, rather than by a single residual sector (Ulyssea, 2020). After a labour-supply shock, informal work may operate as a 'stepping-stone' towards formality in the short run while preserving low-productivity firms over longer horizons (Imbert and Ulyssea, 2026, p. 914). One national household number cannot summarise those dynamics."
)

add_para(
    "For wage-setting, the Kaitz index offers a more relevant starting point. It measures the legal minimum against the median wage of the workers whose employment relationships it regulates. Coverage, compliance and labour-market composition still matter, so the ratio is not a policy rule. It is nevertheless connected to the wage distribution in a way that the DIEESE family budget is not (Dube and Lindner, 2024; OECD, 2026)."
)
add_equation(r"K_t = \frac{MW_t}{\widetilde{W}_t}")
add_para(
    "Using median income from all jobs among employed people in the PNAD Contínua gives a broad Brazilian Kaitz ratio of 73.1 per cent in 2023. The denominator differs from the OECD measure, which uses the gross median wage of full-time employees, but the ratio places the legal floor inside the observed earnings distribution. The full DIEESE family estimate yields a pseudo-Kaitz ratio above 350 per cent, a number driven chiefly by comparing a family budget with individual labour income (IBGE, 2026a; OECD, 2026)."
)

add_figure(12, "Brazil's minimum-to-median ratios, 2012-2023.", "author's calculations from PNAD Contínua, DIEESE and Ipea. The Brazilian denominator includes all occupied workers and all jobs (IBGE, 2026a; DIEESE, 2026; Ipea, 2026).")
add_para(
    "The statutory ratio remains within a comparatively narrow band because the wage floor and median earnings share the broader wage cycle. The per-person DIEESE ratio rose much more after 2020 as food prices moved through the fixed coefficient. That is the movement the formula is designed to capture. It should not be read as evidence that a sustainable legal wage suddenly moved to several times median earnings."
)

add_figure(13, "Minimum-to-median wage ratios in Brazil and OECD economies.", "OECD minimum-to-median ratios and PNAD-based calculations for Brazil. Denominators are not fully harmonised (IBGE, 2026a; OECD, 2026).")
add_para(
    "The international comparison confirms the question of scale. With the broader PNAD denominator, Brazil's legal minimum lies near the upper end of the OECD distribution. The full DIEESE estimate falls outside the relevant range because it was never designed as an individual wage statistic. The figure is useful only if that difference in units is kept in view (IBGE, 2026a; OECD, 2026)."
)

add_para(
    "Purchasing-power parity answers a different question. PPP conversion factors translate currencies using relative price levels and permit comparisons across countries. They are not local cost-of-living indices and cannot identify the budget needed by a family in São Luís or São Paulo. They can, however, compare the purchasing power of Brazil's minimum wage with international poverty thresholds (World Bank, 2025; OECD, 2025)."
)
add_para(
    "A historical PPP series cannot be constructed by dividing every past nominal wage by today's conversion rate. The benchmark must be extrapolated with the difference between Brazilian and American inflation. The standard principle is:"
)
add_equation(r"PPP_t = PPP_b \times \frac{CPI_{BR,t}/CPI_{BR,b}}{CPI_{US,t}/CPI_{US,b}}")
add_para(
    "The Ipea series begins from a 2011 PPP benchmark and extrapolates it with relative consumer-price inflation. The historical real-wage series is then rebased to the 2021 PPP level so that the wage and the World Bank poverty lines are expressed in the same international dollars (Ipea, 2026; World Bank, 2025)."
)
add_equation(r"MW_t^{2021\,int\$} = MW_t^{real,Jun2026\,BRL} \times \frac{MW_{2021}^{PPP,current}}{MW_{2021}^{real,Jun2026\,BRL}}")

add_figure(14, "Brazil's minimum wage and World Bank poverty lines, 1940-2026.", "author's calculations from Ipea and World Bank series. Poverty lines are expressed per person in constant 2021 international dollars (Ipea, 2026; World Bank, 2025).")
add_para(
    "Over most of the period, one statutory minimum wage was worth more than one upper-middle-income poverty threshold for a single person, though the margin varied widely. That comparison describes purchasing power. It does not estimate the income required by a four-person household, and the World Bank lines were never intended to serve as middle-class family budgets (World Bank, 2025)."
)

add_heading('A local reference budget')

add_para(
    "A better measure of household need begins by changing the object being measured. The primary statistic should be a local family budget. A wage comparison can be derived afterwards, once the household type and the number of earners are stated. The Local Adequate Minimum Budget, or OMAL, is an experimental attempt to build that statistic from the 2017-18 POF, quarterly PNAD Contínua income data and the nine local expenditure groups of the IPCA (IBGE, 2019, 2023, 2026a, 2026b)."
)
add_para(
    "The exercise covers sixteen urban areas for which a local IPCA series is available: Rio Branco, São Luís, Aracaju, Campo Grande, Goiânia, Brasília, Belém, Fortaleza, Recife, Salvador, Belo Horizonte, Grande Vitória, Rio de Janeiro, São Paulo, Curitiba and Porto Alegre. It reports six household types and two versions of the budget. The core basket is restricted to recurring expenditures. The expanded basket adds household articles, personal services, leisure and occasional travel, while continuing to exclude vehicle purchases."
)

add_heading('How OMAL is constructed', level=2)

add_para(
    "The first step is to make household size explicit. The current version uses the modified OECD scale: one for the first adult, 0.5 for each additional adult and 0.3 for each child under 14. If A_h is the number of adults and C_h the number of children in household h, adult equivalents are defined as (OECD, 2013):"
)
add_equation(r"AE_h = 1 + 0.5(A_h-1) + 0.3C_h")
add_para(
    "A household with two adults and two children therefore has 2.1 adult equivalents. Disposable income is divided by this scale before households are assigned to the reference band:"
)
add_equation(r"y_h^{eq} = \frac{Y_h^{disp}}{AE_h}")
add_para(
    "The reference sample is restricted to households recorded as food secure, without arrears on rent, mortgage or basic utility bills, and between the 30th and 60th percentiles of equivalised disposable income. In the processed POF data, those cutoffs are R$2,308.64 and R$4,543.71 in the survey's price base. Of 9,844 usable consumption units, 1,481 satisfy the reference conditions (IBGE, 2019, 2023). In set notation:"
)
add_equation(r"\mathcal{R}=\{h:FS_h=1,\;ARR_h=0,\;q_{0.30}\le y_h^{eq}\le q_{0.60}\}")
add_para(
    "This choice does not claim that middle-income consumption automatically defines adequacy. It avoids two more obvious errors. Using the poorest observed households would convert deprivation into a needs standard; using the mean of the whole distribution would import the discretionary consumption of affluent households. The reference band remains normative and should be tested against alternative cutoffs."
)
add_para(
    "POF expenses are converted to monthly values according to each item's reference period and assigned to one of the nine IPCA groups. For expenditure item j in group g, with value x_hj and a reference period of m_j months, the monthly group expenditure is:"
)
add_equation(r"X_{hg}^{m}=\sum_{j\in g}\frac{x_{hj}}{m_j}")
add_para(
    "Food enters group 1; rent, utilities and housing maintenance enter group 2; household articles enter group 3; clothing group 4; transport group 5; health and personal care group 6; recreation and personal services group 7; education group 8; and telephone, mobile and internet services group 9. Imputed rent is retained to measure the housing service consumed by owner-occupiers. Other non-monetary services from the POF2 and POF4 files are not included in the current version, a limitation discussed below (IBGE, 2023)."
)
add_para(
    "Within the reference sample, the estimator uses the survey-weighted median of expenditure per adult equivalent. A median is less sensitive than a mean to a small number of very large purchases. For city c, group g and budget version v, the local estimate is the value that minimises weighted absolute deviations:"
)
add_equation(r"\widehat b_{cgv}=\arg\min_b\sum_{h\in\mathcal{R}_c}w_h\left|\frac{X_{hgv}^{m}}{AE_h}-b\right|")
add_para(
    "Some city samples are small. A nominal count also exaggerates precision when survey weights are unequal, so the diagnostic uses Kish's effective sample size (Kish, 1965):"
)
add_equation(r"n_{eff,c}=\frac{\left(\sum_h w_h\right)^2}{\sum_h w_h^2}")
add_para(
    "The city estimate is then shrunk towards its regional estimate. The amount of shrinkage depends on effective sample size, with a tuning constant of 80:"
)
add_equation(r"\widetilde b_{cgv}=\lambda_c\widehat b_{cgv}+(1-\lambda_c)\widehat b_{rgv},\qquad \lambda_c=\frac{n_{eff,c}}{n_{eff,c}+80}")
add_para(
    "This is an empirical-Bayes-style stabilisation rather than a full hierarchical model. A city with a large effective sample retains most of its own estimate; a city with few effective observations is pulled towards the regional median. Belém, for example, has an effective reference sample of about 12.8 households, while São Paulo has about 121.3. The published spreadsheet reports both the raw and effective sample sizes so that users can see where the estimate relies more heavily on regional information (IBGE, 2019)."
)
add_para(
    "The POF expenditure base is carried to December 2019 using the national IPCA rates of 3.75 per cent in 2018 and 4.31 per cent in 2019. The bridge factor is therefore:"
)
add_equation(r"B_{2019m12}=(1+0.0375)(1+0.0431)=1.08221625")
add_para(
    "From January 2020 onwards, each expenditure group is updated with its own local IPCA series. If π_cgτ is monthly inflation for city c, group g and month τ, the cumulative price index is:"
)
add_equation(r"P_{cgt}=\prod_{\tau=2020m1}^{t}\left(1+\frac{\pi_{cg\tau}}{100}\right)")
add_para(
    "For household type h and budget version v, the monthly OMAL is the sum of the nine updated components:"
)
add_equation(r"OMAL_{cht}^{v}=AE_h\sum_{g=1}^{9}\widetilde b_{cgv}\times B_{2019m12}\times P_{cgt}")
add_para(
    "The formula does not use PNAD earnings to determine the budget. Doing so would make poor cities appear to need less merely because their residents earn less. PNAD is retained as an external comparison. The budget can also be translated into required net income per worker when the number of earners n is stated:"
)
add_equation(r"W_{cht}^{net}(n)=\frac{OMAL_{cht}^{v}}{n}")
add_para(
    "Two ratios are reported for interpretation. The first compares required income per worker with the legal minimum. The second compares the family budget with the DIEESE family benchmark:"
)
add_equation(r"R_{cht}^{MW}(n)=\frac{W_{cht}^{net}(n)}{MW_t},\qquad R_{cht}^{DIEESE}=\frac{OMAL_{cht}^{v}}{SMN_t}")
add_para(
    "These are diagnostic ratios, not policy targets. Gross earnings would require a further inversion of payroll taxes and income-tax schedules, which the present version does not attempt. Nor does the model estimate confidence intervals. The effective sample size and shrinkage weight provide a warning about precision, but a publishable statistical release should add bootstrap or replicate-weight intervals."
)

add_heading('What the estimates show', level=2)

add_figure(15, "OMAL by city, June 2026.", "author's calculations from POF, local IPCA and PNAD files (IBGE, 2019, 2023, 2026a, 2026b).")
add_para(
    "For a household with two adults and two children, the core OMAL ranges from R$5,976.93 in Grande Vitória to R$6,969.21 in São Paulo in June 2026. The expanded basket ranges from R$6,328.24 to R$7,264.66. The ordering changes little when discretionary but common expenditures are added. The differences are therefore driven mainly by the recurring cost of maintaining a household rather than by leisure or durable purchases (IBGE, 2019, 2026b)."
)

add_figure(16, "Composition of the core OMAL basket, June 2026.", "author's calculations from POF and local IPCA files (IBGE, 2019, 2023, 2026b).")
add_para(
    "Housing accounts for about 37 per cent of the average core budget across the sixteen cities, food for 25 per cent, health and personal care for 14 per cent and transport for 12 per cent. Education, clothing and communication make up the remainder. This composition explains why the nine-group update is preferable even though it changes the final total only modestly. The model can show which prices are moving the budget rather than hiding them inside an aggregate index."
)

add_figure(17, "Core OMAL over time in selected cities, 2020-2026.", "author's calculations from POF and local IPCA files (IBGE, 2019, 2026b).")
add_para(
    "The time series shows a common inflationary cycle and persistent differences in levels. Costs accelerated after 2020, especially between 2021 and 2023, and then continued to rise more gradually. São Paulo remains the most expensive city in the selected series; Grande Vitória stays near the lower edge. The persistence is useful for policy because it separates durable geographic differences from temporary monthly noise (IBGE, 2026b)."
)

add_figure(18, "Mean labour income and the core OMAL.", "OMAL for June 2026 and mean real labour income for 2026Q1. The diagonal marks equality (IBGE, 2019, 2026a, 2026b).")
add_para(
    "The family budget is closest to mean labour income in Brasília, where the ratio is about 0.91. In São Luís the family budget is slightly more than twice average labour income. That contrast should not be read as a city-specific wage prescription. It shows why the units must be kept separate. A household budget above one worker's earnings may be sustained by two earners, transfers or other income. The relevant policy question is not whether one worker can finance every family arrangement, but how the legal floor interacts with household composition and local labour markets (IBGE, 2026a)."
)

add_figure(19, "Effect of using all nine local IPCA groups.", "comparison with the earlier simplified update rule (IBGE, 2026b).")
add_para(
    "Replacing the simplified update with all nine IPCA groups reduces the core OMAL by about 0.76 per cent on average. The largest absolute change appears in Goiânia. The small aggregate revision is reassuring rather than trivial. The earlier approximation was not wildly wrong, but the full specification is more coherent, especially for decomposition and monthly change. Measurement improves when the internal structure is correct even if the headline total barely moves."
)

add_figure(20, "Core OMAL as a share of the DIEESE family benchmark, June 2026.", "author's calculations from OMAL and DIEESE data (DIEESE, 2026; IBGE, 2019, 2026b).")
add_para(
    "The core OMAL lies below the DIEESE family estimate in every city, ranging from 73.7 per cent of the benchmark in Grande Vitória to 85.9 per cent in São Paulo. The difference reflects more than geography. DIEESE scales an old food basket through a fixed food share; OMAL estimates the main spending categories directly and updates each with a local price index. The two statistics answer related but distinct questions."
)

add_figure(21, "Average core OMAL by household type across the sixteen cities.", "author's calculations from POF and local IPCA files (IBGE, 2019, 2023, 2026b).")
add_para(
    "Household composition changes the result almost as much as geography. The average budget rises from one adult to lone-parent households and then to two-adult families with children. The increments follow the stated equivalence scale rather than the assumption that every additional person adds the same cost. This feature makes the model open to revision: a different scale can be substituted and the entire set of estimates recalculated (OECD, 2013)."
)

add_figure(22, "Per-worker OMAL under a two-worker assumption, mean labour income and the statutory minimum wage, June 2026.", "author's calculations from OMAL, PNAD and Ipea data (IBGE, 2019, 2026a, 2026b; Ipea, 2026).")
add_para(
    "When the two-adult, two-child budget is divided between two workers, required net income ranges from 1.84 to 2.15 statutory minimum wages per worker. The result is much closer to observed average earnings than the full family benchmark is, though it remains above the legal minimum in every city. This is the comparison that wage policy needs: household costs translated through an explicit assumption about earners, then placed beside the wage distribution and the legal floor."
)

add_heading('What the number can do')

add_para(
    "OMAL is not a new official line of need. Its reference group, equivalence scale, expenditure boundary and shrinkage parameter all contain judgments. The current version also omits most non-monetary services, uses group-level rather than subitem-level IPCA indices and reports no sampling interval. Those limitations should be treated as part of the statistic, not hidden in a technical appendix (IBGE, 2023)."
)
add_para(
    "Even with those qualifications, the framework is better suited to policy analysis than the DIEESE formula. It reports separate household types, prices the main non-food categories directly, uses local inflation and distinguishes a family budget from an individual wage. It also preserves the political insight behind the DIEESE series: workers need an intelligible account of the cost of living that is not supplied solely by employers or government."
)
add_para(
    "The two measures can therefore coexist. DIEESE can continue to publish a demanding bargaining benchmark rooted in trade-union advocacy. A public statistical system can publish local reference budgets with explicit assumptions and uncertainty. Minimum-wage policy should then be informed by those budgets alongside the Kaitz ratio, compliance, regional and sectoral bite, employment, hours, informal spillovers, prices, productivity and fiscal exposure (Dube and Lindner, 2024; Derenoncourt et al., 2025; Engbom and Moser, 2022; Instituição Fiscal Independente, 2026)."
)
add_para(
    "The necessary minimum wage is useful in a narrower role than its public reputation suggests. It keeps the cost of living in view and gives workers a common number in negotiations. It does not identify the needs of a representative Brazilian household and should not organise wage-setting or anti-poverty policy. A local reference budget such as OMAL does not remove judgment from the exercise. It makes the judgment visible, disaggregated and revisable. That is a more credible basis for policy."
)

# ---------------- References ----------------
add_heading('References')
references = [
    "Brochu, P. R., Green, D. A., Lemieux, T., & Townsend, J. H. (2025). The minimum wage, turnover, and the shape of the wage distribution. NBER Working Paper 33479.",
    "Brynjolfsson, E., Li, J. F., Miranda, J., Seamans, R., & Wang, A. J. (2026). Minimum wages and the rise of robots. NBER Working Paper 34895.",
    "Davies, M., Park, R. J., & Stansbury, A. (2026). Minimum wages and workplace injuries. Working paper.",
    "Derenoncourt, E., Gerard, F., Lagos, L., & Montialoux, C. (2025). Minimum wages and informality. NBER Working Paper 34445.",
    "DIEESE. (2016). Metodologia da Pesquisa Nacional da Cesta Básica de Alimentos. São Paulo: Departamento Intersindical de Estatística e Estudos Socioeconômicos.",
    "DIEESE. (2025). Quem somos. São Paulo: Departamento Intersindical de Estatística e Estudos Socioeconômicos.",
    "DIEESE. (2026). Salário mínimo nominal e necessário and Pesquisa Nacional da Cesta Básica de Alimentos, monthly series through June 2026.",
    "Dube, A., & Lindner, A. S. (2024). Minimum wages in the 21st century. NBER Working Paper 32878.",
    "Engbom, N., & Moser, C. (2022). Earnings inequality and the minimum wage: Evidence from Brazil. American Economic Review, 112(12), 3803-3847.",
    "IBGE. (2004). Pesquisa de Orçamentos Familiares 2002-2003. Rio de Janeiro: Instituto Brasileiro de Geografia e Estatística.",
    "IBGE. (2010). Pesquisa de Orçamentos Familiares 2008-2009. Rio de Janeiro: Instituto Brasileiro de Geografia e Estatística.",
    "IBGE. (2019). Pesquisa de Orçamentos Familiares 2017-2018: primeiros resultados. Rio de Janeiro: Instituto Brasileiro de Geografia e Estatística.",
    "IBGE. (2023). Microdados da Pesquisa de Orçamentos Familiares 2017-2018: leia-me, arquivos de dados, tradutores and memórias de cálculo. Rio de Janeiro: Instituto Brasileiro de Geografia e Estatística.",
    "IBGE. (2026a). PNAD Contínua trimestral: rendimento médio mensal habitual de todos os trabalhos. SIDRA table 6472.",
    "IBGE. (2026b). IPCA by area and expenditure group. SIDRA table 7060.",
    "IBGE. (2026c). Taxa de informalidade das pessoas de 14 anos ou mais ocupadas. SIDRA table 8529.",
    "Imbert, C., & Ulyssea, G. (2026). Rural migrants and urban informality: Evidence from Brazil. Econometrica, 94(3), 911-939.",
    "Instituição Fiscal Independente. (2026). Relatório de Acompanhamento Fiscal 111. Brasília: Senado Federal.",
    "Ipea. (2026). Salário mínimo real (GAC12_SALMINRE12) and salário mínimo em paridade do poder de compra (GAC12_SALMINDOL12). Ipeadata.",
    "Jardim, E., Long, M. C., Plotnick, R., van Inwegen, E., Vigdor, J., & Wething, H. (2018). Minimum wage increases and individual employment trajectories. NBER Working Paper 25182.",
    "Kish, L. (1965). Survey Sampling. New York: Wiley.",
    "Melo, V. C., Kaiser, C., Neumark, D., Palagashvili, L., & Farren, M. D. (2026). Minimum wage laws and job search. NBER Working Paper 33433, revised July 2026.",
    "OECD. (2013). OECD Framework for Statistics on the Distribution of Household Income, Consumption and Wealth. Paris: OECD Publishing.",
    "OECD. (2025). Purchasing power parities database. Paris: Organisation for Economic Co-operation and Development.",
    "OECD. (2026). Employment Outlook 2026. Paris: Organisation for Economic Co-operation and Development.",
    "Pires, M. (2025). Repressão salarial: os dilemas do salário mínimo. Blog do IBRE, 5 June.",
    "Soares, S. S. D. (2009). Metodologias para estabelecer a linha de pobreza: objetivas, subjetivas, relativas, multidimensionais. Texto para Discussão 1381. Brasília: Ipea.",
    "Stansbury, A., & Summers, L. H. (2020). The declining worker power hypothesis: An explanation for the recent evolution of the American economy. Brookings Papers on Economic Activity, Spring, 1-96.",
    "Ulyssea, G. (2020). Informality: Causes and consequences for development. Annual Review of Economics, 12, 525-546.",
    "World Bank. (2025). Poverty and Inequality Platform and methodology for extrapolating PPP conversion factors. Washington, DC: World Bank.",
]
for ref in references:
    p = doc.add_paragraph(style='Reference')
    p.add_run(ref)

# Save
# Core properties
doc.core_properties.title = 'Is the DIEESE minimum wage useful for policy?'
doc.core_properties.subject = 'DIEESE necessary minimum wage and OMAL local reference budget'
doc.core_properties.author = 'João Gabriel Caetano Leite'
doc.save(OUT)

# ---------------- Markdown companion ----------------
# Create a readable markdown version by traversing document blocks.
md_lines = [
    '# Is the DIEESE minimum wage useful for policy?',
    '',
    '*DIEESE gave Brazilian workers an independent economic vocabulary. Its “necessary minimum wage” remains useful in bargaining, but the formula is too rigid to measure household need and too detached from labour-market evidence to guide the statutory wage floor.*',
    '',
    'The complete editorial revision, formulas, citations and figures are contained in the DOCX. The markdown companion reproduces the OMAL construction formulas and links the final figures.',
    '',
    '## OMAL formulas',
    '',
    r'$$AE_h = 1 + 0.5(A_h-1) + 0.3C_h$$',
    '',
    r'$$y_h^{eq} = \frac{Y_h^{disp}}{AE_h}$$',
    '',
    r'$$\mathcal{R}=\{h:FS_h=1,\;ARR_h=0,\;q_{0.30}\le y_h^{eq}\le q_{0.60}\}$$',
    '',
    r'$$X_{hg}^{m}=\sum_{j\in g}\frac{x_{hj}}{m_j}$$',
    '',
    r'$$\widehat b_{cgv}=\arg\min_b\sum_{h\in\mathcal{R}_c}w_h\left|\frac{X_{hgv}^{m}}{AE_h}-b\right|$$',
    '',
    r'$$n_{eff,c}=\frac{\left(\sum_h w_h\right)^2}{\sum_h w_h^2}$$',
    '',
    r'$$\widetilde b_{cgv}=\lambda_c\widehat b_{cgv}+(1-\lambda_c)\widehat b_{rgv},\qquad \lambda_c=\frac{n_{eff,c}}{n_{eff,c}+80}$$',
    '',
    r'$$P_{cgt}=\prod_{\tau=2020m1}^{t}\left(1+\frac{\pi_{cg\tau}}{100}\right)$$',
    '',
    r'$$OMAL_{cht}^{v}=AE_h\sum_{g=1}^{9}\widetilde b_{cgv}\times 1.08221625\times P_{cgt}$$',
    '',
    r'$$W_{cht}^{net}(n)=\frac{OMAL_{cht}^{v}}{n}$$',
    '',
    r'$$R_{cht}^{MW}(n)=\frac{W_{cht}^{net}(n)}{MW_t},\qquad R_{cht}^{DIEESE}=\frac{OMAL_{cht}^{v}}{SMN_t}$$',
    '',
    '## OMAL figures',
    ''
]
for n in range(15,23):
    md_lines.append(f'![Figure {n}](figures/{figures[n].name})')
    md_lines.append('')
MD_OUT.write_text('\n'.join(md_lines), encoding='utf-8')

# ---------------- Package ----------------
if PKG.exists():
    shutil.rmtree(PKG)
PKG.mkdir()
(PKG/'figures').mkdir()
shutil.copy2(OUT, PKG/OUT.name)
shutil.copy2(MD_OUT, PKG/MD_OUT.name)
spreadsheet = TABLE_DIR/'OMAL_indicador_9_grupos_IPCA_2020_2026.xlsx'
if spreadsheet.exists():
    shutil.copy2(spreadsheet, PKG/spreadsheet.name)
for n,p in figures.items():
    shutil.copy2(p, PKG/'figures'/f'figure{n:02d}_{p.name}')
with ZipFile(ZIP_OUT, 'w', compression=ZIP_DEFLATED) as z:
    for p in sorted(PKG.rglob('*')):
        if p.is_file():
            z.write(p, arcname=str(p.relative_to(PKG.parent)))

print(OUT)
print(MD_OUT)
print(ZIP_OUT)
